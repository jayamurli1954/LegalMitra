"""
Document Processor Service
Handles extraction of text from various document formats
"""

import io
import sys
import os
from typing import Optional

# PDF processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDF_PLUMBER_AVAILABLE = True
except ImportError:
    PDF_PLUMBER_AVAILABLE = False

# Word document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Image processing and OCR
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

# PDF to image conversion (for image-based PDFs)
try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# AI Vision APIs (for better image processing)
# Note: We now use the new google.genai package (not deprecated google.generativeai)
# The gemini_ocr utility handles the actual API calls using the new SDK

try:
    from openai import OpenAI
    OPENAI_VISION_AVAILABLE = True
except Exception:
    OPENAI_VISION_AVAILABLE = False

from app.core.config import get_settings
import base64
import logging


class DocumentProcessor:
    """Service for processing and extracting text from various document formats"""
    
    def __init__(self):
        self.settings = get_settings()
    
    async def process_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        logger = logging.getLogger(__name__)
        
        if not PDF_AVAILABLE and not PDF_PLUMBER_AVAILABLE:
            raise Exception("PDF processing libraries not available. Please install PyPDF2 or pdfplumber.")
        
        pdf_file = io.BytesIO(file_content)
        text_parts = []
        last_error = None
        
        # Try pdfplumber first (better text extraction)
        if PDF_PLUMBER_AVAILABLE:
            try:
                pdf_file.seek(0)  # Reset file pointer
                with pdfplumber.open(pdf_file) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(page_text)
                        except Exception as page_error:
                            logger.warning(f"Error extracting text from page {page_num} with pdfplumber: {page_error}")
                            continue
                
                if text_parts:
                    extracted_text = "\n\n".join(text_parts)
                    # Check if extracted text is meaningful (not just metadata/whitespace)
                    # If text is too short (< 50 chars), it might be just metadata, try OCR
                    if extracted_text.strip() and len(extracted_text.strip()) >= 50:
                        return extracted_text
                    else:
                        logger.info(f"Extracted text too short ({len(extracted_text.strip())} chars), likely metadata. Will try OCR.")
                        text_parts = []  # Clear to trigger OCR
            except Exception as e:
                last_error = f"pdfplumber error: {str(e)}"
                logger.warning(f"pdfplumber failed: {last_error}")
        
        # Fallback to PyPDF2 if pdfplumber failed or returned no text
        if PDF_AVAILABLE and not text_parts:
            try:
                pdf_file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            text_parts.append(text)
                    except Exception as page_error:
                        logger.warning(f"Error extracting text from page {page_num} with PyPDF2: {page_error}")
                        continue
                
                if text_parts:
                    extracted_text = "\n\n".join(text_parts)
                    # Check if extracted text is meaningful (not just metadata/whitespace)
                    # If text is too short (< 50 chars), it might be just metadata, try OCR
                    if extracted_text.strip() and len(extracted_text.strip()) >= 50:
                        return extracted_text
                    else:
                        logger.info(f"Extracted text too short ({len(extracted_text.strip())} chars), likely metadata. Will try OCR.")
                        text_parts = []  # Clear to trigger OCR
            except Exception as e:
                last_error = f"PyPDF2 error: {str(e)}"
                logger.warning(f"PyPDF2 failed: {last_error}")
        
        # If no text was extracted, try OCR on image-based PDF
        # Convert PDF pages to images and use Gemini OCR
        ocr_error_msg = None  # Store OCR error if it occurs
        ocr_attempted = False  # Track if OCR was attempted
        if not text_parts and PDF2IMAGE_AVAILABLE and PIL_AVAILABLE:
            ocr_attempted = True
            try:
                logger.info("No text extracted from PDF. Attempting OCR on image-based PDF...")
                pdf_file.seek(0)
                
                # Check if Poppler is available (required for pdf2image on Windows)
                try:
                    import subprocess
                    import shutil
                    
                    # First, try to find pdftoppm in PATH
                    pdftoppm_path = shutil.which('pdftoppm')
                    
                    if not pdftoppm_path:
                        # Try common Windows installation paths
                        common_paths = [
                            r"C:\Program Files\poppler\Library\bin\pdftoppm.exe",
                            r"C:\poppler\Library\bin\pdftoppm.exe",
                            r"C:\Program Files (x86)\poppler\Library\bin\pdftoppm.exe",
                        ]
                        
                        for path in common_paths:
                            if os.path.exists(path):
                                pdftoppm_path = path
                                # Temporarily add to PATH for this process
                                bin_dir = os.path.dirname(path)
                                os.environ['PATH'] = bin_dir + os.pathsep + os.environ.get('PATH', '')
                                logger.info(f"Found Poppler at: {bin_dir}, added to PATH for this session")
                                break
                    
                    if pdftoppm_path:
                        # Test if it works
                        result = subprocess.run(
                            [pdftoppm_path, '-h'], 
                            capture_output=True, 
                            text=True, 
                            timeout=2,
                            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
                        )
                        if result.returncode != 0 and 'pdftoppm' not in result.stdout and 'pdftoppm' not in result.stderr:
                            raise FileNotFoundError("pdftoppm command not working correctly")
                    else:
                        raise FileNotFoundError("pdftoppm command not found")
                        
                except (FileNotFoundError, subprocess.TimeoutExpired, OSError) as e:
                    ocr_error_msg = f"Poppler is not installed or not in PATH. pdf2image requires Poppler to convert PDF pages to images. Error: {str(e)}"
                    logger.error(ocr_error_msg)
                    
                    # Check if Poppler exists in common paths but not in PATH
                    common_bin_paths = [
                        r"C:\Program Files\poppler\Library\bin",
                        r"C:\poppler\Library\bin",
                    ]
                    
                    found_path = None
                    for path in common_bin_paths:
                        if os.path.exists(os.path.join(path, "pdftoppm.exe")):
                            found_path = path
                            break
                    
                    if found_path:
                        ocr_error_msg += f"\n\n[SOLUTION] Poppler is installed at: {found_path}\n"
                        ocr_error_msg += "But it's not in your system PATH. Restart your server after adding it to PATH:\n"
                        ocr_error_msg += f"1. Add this to PATH: {found_path}\n"
                        ocr_error_msg += "2. Restart your server/terminal for PATH changes to take effect."
                    
                    raise Exception(ocr_error_msg)
                
                # Convert PDF pages to images
                images = convert_from_bytes(file_content, dpi=300)  # Higher DPI for better OCR quality
                logger.info(f"Converted PDF to {len(images)} page images")
                
                # Extract text from each page image
                # Try Gemini OCR first, fallback to OpenAI Vision if available
                page_texts = []
                use_openai_fallback = False
                
                # Try Gemini OCR first
                try:
                    from app.utils.gemini_ocr import extract_text_from_image
                    gemini_available = True
                except ImportError:
                    gemini_available = False
                    logger.warning("Gemini OCR not available, will try OpenAI Vision if configured")
                
                for page_num, image in enumerate(images, 1):
                    try:
                        # Convert PIL Image to bytes
                        img_byte_arr = io.BytesIO()
                        image.save(img_byte_arr, format='PNG')
                        img_bytes = img_byte_arr.getvalue()
                        
                        page_text = None
                        
                        # Try Gemini OCR first (if available and not already failed)
                        if gemini_available and not use_openai_fallback:
                            try:
                                logger.info(f"Processing page {page_num}/{len(images)} with Gemini OCR...")
                                page_text = extract_text_from_image(img_bytes, mime_type="image/png")
                                if page_text and page_text.strip():
                                    page_texts.append(f"--- Page {page_num} ---\n{page_text}")
                                    logger.info(f"Extracted {len(page_text)} characters from page {page_num} using Gemini")
                                    continue
                            except Exception as gemini_error:
                                error_str = str(gemini_error).lower()
                                # If Gemini fails due to API key issues, try OpenAI fallback
                                if "leaked" in error_str or "permission_denied" in error_str or "403" in str(gemini_error) or "authentication" in error_str:
                                    logger.warning(f"Gemini OCR failed on page {page_num}, trying OpenAI Vision fallback...")
                                    use_openai_fallback = True
                                    # Store error but continue to OpenAI fallback
                                    if not ocr_error_msg:
                                        ocr_error_msg = f"Gemini OCR failed: {str(gemini_error)}. Trying OpenAI Vision as fallback."
                                else:
                                    # Other Gemini errors - store and continue
                                    if not ocr_error_msg:
                                        ocr_error_msg = f"OCR error on page {page_num}: {str(gemini_error)}"
                        
                        # Fallback to OpenAI Vision if Gemini failed or not available
                        if (use_openai_fallback or not gemini_available) and OPENAI_VISION_AVAILABLE and self.settings and self.settings.OPENAI_API_KEY:
                            try:
                                logger.info(f"Processing page {page_num}/{len(images)} with OpenAI Vision...")
                                client = OpenAI(api_key=self.settings.OPENAI_API_KEY)
                                image_base64 = base64.b64encode(img_bytes).decode('utf-8')
                                
                                response = client.chat.completions.create(
                                    model="gpt-4o",
                                    messages=[
                                        {
                                            "role": "user",
                                            "content": [
                                                {"type": "text", "text": "Extract all text from this legal document image. Preserve formatting and structure."},
                                                {
                                                    "type": "image_url",
                                                    "image_url": {
                                                        "url": f"data:image/png;base64,{image_base64}"
                                                    }
                                                }
                                            ]
                                        }
                                    ],
                                    max_tokens=4000
                                )
                                
                                if response and response.choices and response.choices[0].message.content:
                                    page_text = response.choices[0].message.content
                                    if page_text and page_text.strip():
                                        page_texts.append(f"--- Page {page_num} ---\n{page_text}")
                                        logger.info(f"Extracted {len(page_text)} characters from page {page_num} using OpenAI Vision")
                                        continue
                            except Exception as openai_error:
                                logger.error(f"OpenAI Vision also failed on page {page_num}: {openai_error}")
                                if not ocr_error_msg or "Gemini" not in ocr_error_msg:
                                    ocr_error_msg = f"Both Gemini and OpenAI Vision failed. Last error: {str(openai_error)}"
                        
                        # If no text extracted from this page
                        if not page_text or not page_text.strip():
                            logger.warning(f"No text extracted from page {page_num}")
                            
                    except Exception as page_error:
                        logger.error(f"Error processing page {page_num} with OCR: {page_error}")
                        if not ocr_error_msg:
                            ocr_error_msg = f"OCR error on page {page_num}: {str(page_error)}"
                        continue
                
                if page_texts:
                    extracted_text = "\n\n".join(page_texts)
                    logger.info(f"Successfully extracted text from {len(page_texts)} pages using OCR")
                    return extracted_text
                else:
                    logger.warning("OCR processing completed but no text was extracted from any page")
                    # If no text was extracted but no errors were raised, it might be an API issue
                    if not ocr_error_msg:
                        ocr_error_msg = "OCR processing completed but returned no text from any page. This could indicate an API key issue, empty pages, or API quota limit."
                    
            except Exception as ocr_error:
                ocr_error_msg = str(ocr_error)
                logger.error(f"OCR fallback failed: {ocr_error_msg}")
                # Include full traceback in logs for debugging
                import traceback
                logger.error(f"OCR error traceback: {traceback.format_exc()}")
        
        # If OCR also failed or is not available, raise an informative error
        error_msg = "Could not extract text from PDF. "
        if last_error:
            error_msg += f"Last error: {last_error}. "
        
        if not PDF2IMAGE_AVAILABLE:
            error_msg += "The PDF appears to be image-based (scanned). "
            error_msg += "Install pdf2image package (pip install pdf2image) to enable OCR processing of scanned PDFs. "
            error_msg += "Note: On Windows, you may also need to install poppler (https://github.com/oschwartz10612/poppler-windows/releases)."
        else:
            error_msg += "The PDF might be image-based (scanned) or encrypted"
            if ocr_attempted:
                error_msg += ", and OCR processing also failed"
            else:
                error_msg += ". OCR processing was not attempted"
            error_msg += ". "
            
            # Include detailed OCR error information
            if ocr_error_msg:
                error_msg += f"\n\n**OCR Error Details:** {ocr_error_msg}\n\n"
                
                # Provide specific guidance based on error type
                if "leaked" in ocr_error_msg.lower():
                    error_msg += "**⚠️ SECURITY ALERT: Your Gemini API key has been reported as leaked.**\n\n"
                    error_msg += "**IMMEDIATE ACTION REQUIRED:**\n"
                    error_msg += "1. **Revoke the current API key** in Google AI Studio: https://makersuite.google.com/app/apikey\n"
                    error_msg += "2. **Generate a new API key** from the same page\n"
                    error_msg += "3. **Update your .env file** with the new key: GOOGLE_GEMINI_API_KEY=your_new_key\n"
                    error_msg += "4. **Restart your server** for changes to take effect\n\n"
                    error_msg += "**Alternative OCR Options:**\n"
                    error_msg += "- If you have OPENAI_API_KEY configured, the system will automatically use OpenAI Vision API\n"
                    error_msg += "- You can also install Tesseract OCR for local OCR processing\n\n"
                elif "api_key" in ocr_error_msg.lower() or "authentication" in ocr_error_msg.lower() or "GOOGLE_GEMINI_API_KEY" in ocr_error_msg or "permission_denied" in ocr_error_msg.lower():
                    error_msg += "**This appears to be an API key configuration issue.**\n"
                    error_msg += "Please ensure GOOGLE_GEMINI_API_KEY is set correctly in your .env file.\n"
                    error_msg += "You can get a Gemini API key from: https://makersuite.google.com/app/apikey\n"
                    error_msg += "**Alternative:** Configure OPENAI_API_KEY to use OpenAI Vision API for OCR instead.\n\n"
                elif "poppler" in ocr_error_msg.lower() or "pdfinfo" in ocr_error_msg.lower() or "pdftoppm" in ocr_error_msg.lower():
                    error_msg += "**This appears to be a Poppler installation issue.**\n"
                    error_msg += "On Windows, pdf2image requires Poppler to be installed.\n"
                    error_msg += "Download from: https://github.com/oschwartz10612/poppler-windows/releases\n"
                    error_msg += "Add the Poppler bin directory to your system PATH.\n\n"
                elif "quota" in ocr_error_msg.lower() or "limit" in ocr_error_msg.lower():
                    error_msg += "**This appears to be a Gemini API quota issue.**\n"
                    error_msg += "Check your Google Cloud Console for API usage and quotas.\n\n"
                elif "404" in ocr_error_msg or "not found" in ocr_error_msg.lower():
                    error_msg += "**This appears to be a Gemini model availability issue.**\n"
                    error_msg += "The Gemini model may not be available. Check Google's API status.\n\n"
                else:
                    error_msg += "**Possible solutions:**\n"
                    error_msg += "1. Verify GOOGLE_GEMINI_API_KEY in .env file\n"
                    error_msg += "2. Check if Poppler is installed (Windows only)\n"
                    error_msg += "3. Check server console logs for more details\n"
                    error_msg += "4. Verify Gemini API is enabled in Google Cloud Console\n"
                    error_msg += "5. Restart the server after changing .env file\n\n"
            else:
                # If OCR was attempted but no error was captured, provide general troubleshooting
                if ocr_attempted:
                    error_msg += "\n\n**OCR Processing Status:** OCR was attempted but no specific error was captured.\n\n"
                    error_msg += "**Possible causes:**\n"
                    error_msg += "1. Gemini API key not configured or invalid (check GOOGLE_GEMINI_API_KEY in .env)\n"
                    error_msg += "2. API quota exceeded (check Google Cloud Console)\n"
                    error_msg += "3. Poppler not installed (Windows only - required for pdf2image)\n"
                    error_msg += "4. Network connectivity issues\n"
                    error_msg += "5. Server needs restart after .env changes\n\n"
                    error_msg += "**Check server console logs** for detailed error information (look for 'OCR error' or 'OCR fallback failed' messages).\n\n"
                    error_msg += "**To get more detailed errors:**\n"
                    error_msg += "- Ensure GOOGLE_GEMINI_API_KEY is set in .env file\n"
                    error_msg += "- Restart the server after making changes\n"
                    error_msg += "- Check the server console output when uploading the PDF\n"
                else:
                    error_msg += "Please ensure GOOGLE_GEMINI_API_KEY is configured for OCR processing."
        
        raise Exception(error_msg)
    
    async def process_word(self, file_content: bytes, file_extension: str) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            raise Exception("Word document processing not available. Please install python-docx.")
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts) if text_parts else "No text could be extracted from the document."
            
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    async def process_image(self, file_content: bytes, filename: str) -> str:
        """Extract text from image using OCR or AI vision"""
        logger = logging.getLogger(__name__)
        gemini_error = None
        openai_error = None
        last_error = None
        
        # Try Gemini OCR utility first (recommended approach)
        try:
            from app.utils.gemini_ocr import extract_text_from_image, get_mime_type_from_filename
            
            logger.info("Attempting Gemini OCR (using gemini_ocr utility)...")
            mime_type = get_mime_type_from_filename(filename)
            extracted_text = extract_text_from_image(file_content, mime_type)
            logger.info(f"Gemini OCR successful: Extracted {len(extracted_text)} characters")
            return extracted_text
        except ImportError:
            logger.warning("Gemini OCR utility not available, trying fallback method")
        except Exception as e:
            gemini_error = str(e)
            last_error = f"Gemini OCR error: {str(e)}"
            logger.error(f"Gemini OCR failed: {last_error}")
            import traceback
            logger.error(f"Error trace: {traceback.format_exc()}")
        
        # Note: No fallback to old deprecated SDK needed - gemini_ocr utility uses the new google.genai SDK
        # If it fails, we'll proceed to OpenAI or Tesseract fallbacks
        
        # Try OpenAI Vision API (most accurate if available)
        if OPENAI_VISION_AVAILABLE and self.settings:
            try:
                if self.settings.OPENAI_API_KEY:
                    client = OpenAI(api_key=self.settings.OPENAI_API_KEY)
                    
                    # Convert image to base64
                    image_base64 = base64.b64encode(file_content).decode('utf-8')
                    
                    # Determine image MIME type from filename
                    image_mime = "image/jpeg"
                    if filename.lower().endswith('.png'):
                        image_mime = "image/png"
                    
                    response = client.chat.completions.create(
                        model="gpt-4o",  # or gpt-4-vision-preview
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": "Extract all text from this image. If this is a legal document, preserve the formatting and structure."},
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:{image_mime};base64,{image_base64}"
                                        }
                                    }
                                ]
                            }
                        ],
                        max_tokens=4000
                    )
                    if response and response.choices and response.choices[0].message.content:
                        return response.choices[0].message.content
            except Exception as e:
                openai_error = str(e)  # Store OpenAI error separately
                if not gemini_error:  # Only overwrite if Gemini didn't fail first
                    last_error = f"OpenAI Vision API error: {str(e)}"
                logger.warning(f"OpenAI vision failed: {str(e)}")
        
        # Fallback to OCR using pytesseract
        if TESSERACT_AVAILABLE and PIL_AVAILABLE:
            try:
                image = Image.open(io.BytesIO(file_content))
                text = pytesseract.image_to_string(image)
                if text and text.strip():
                    return text
                else:
                    raise Exception("OCR returned empty text")
            except Exception as e:
                last_error = f"OCR error: {str(e)}"
                logger.warning(f"OCR failed: {str(e)}")
        
        # If all methods failed, provide helpful error message
        logger.error(f"All image processing methods failed.")
        if gemini_error:
            logger.error(f"Gemini error: {gemini_error}")
        if openai_error:
            logger.error(f"OpenAI error: {openai_error}")
        if last_error:
            logger.error(f"Last error: {last_error}")
        
        error_msg = "Could not extract text from image.\n\n"
        
        # Prioritize showing Gemini error if it exists (most likely to be the issue)
        if gemini_error:
            error_msg += "**Gemini Vision API Error:**\n\n"
            error_msg += f"Gemini Vision API was attempted but failed with the following error:\n\n"
            error_msg += f"`{gemini_error}`\n\n"
            
            # Provide specific guidance based on common errors
            gemini_lower = gemini_error.lower()
            if "invalid" in gemini_lower or "authentication" in gemini_lower or "api_key" in gemini_lower or "unauthorized" in gemini_lower:
                error_msg += "**This appears to be an authentication error.**\n"
                error_msg += "Please verify your GOOGLE_GEMINI_API_KEY in .env file is correct and not expired.\n\n"
            elif "quota" in gemini_lower or "limit" in gemini_lower or "exceeded" in gemini_lower:
                error_msg += "**This appears to be a quota/rate limit error.**\n"
                error_msg += "Check your Google Cloud Console for API usage limits and quotas.\n\n"
            elif "permission" in gemini_lower or "access" in gemini_lower or "forbidden" in gemini_lower:
                error_msg += "**This appears to be a permissions error.**\n"
                error_msg += "Ensure 'Generative Language API' is enabled in your Google Cloud Console.\n\n"
            elif "not found" in gemini_lower or "404" in gemini_lower or "model" in gemini_lower:
                error_msg += "**This appears to be a model availability error.**\n"
                error_msg += "The Gemini model may not be available. Check Google's API status.\n\n"
            else:
                error_msg += "**Troubleshooting steps:**\n"
                error_msg += "1. Verify your GOOGLE_GEMINI_API_KEY in .env file is correct\n"
                error_msg += "2. Check Google Cloud Console for API usage and quotas\n"
                error_msg += "3. Ensure Generative Language API is enabled in Google Cloud Console\n"
                error_msg += "4. Check your internet connection\n\n"
        else:
            # Check if google-genai package is installed
            try:
                from google import genai
                if not (self.settings and self.settings.GOOGLE_GEMINI_API_KEY):
                    error_msg += "google-genai package is installed but API key is not configured.\n"
                    error_msg += "Add GOOGLE_GEMINI_API_KEY to your .env file.\n\n"
                else:
                    error_msg += "google-genai package is available but OCR failed.\n"
                    error_msg += "Check server console logs for details.\n\n"
            except ImportError:
                error_msg += "google-genai package is not installed.\n"
                error_msg += "Install with: pip install google-genai\n\n"
        
        # Mention other fallback options
        if not gemini_error or (gemini_error and "quota" not in gemini_error.lower()):
            error_msg += "**Alternative options:**\n"
            error_msg += "1. Fix the Gemini API key/configuration (recommended)\n"
            error_msg += "2. Use OpenAI API key (OPENAI_API_KEY in .env) for GPT-4 Vision\n"
            error_msg += "3. Install Tesseract OCR (https://github.com/tesseract-ocr/tesseract/wiki)\n"
        
        raise Exception(error_msg)


# Singleton instance
document_processor = DocumentProcessor()

